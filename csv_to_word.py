from docx import Document
from pydrive.auth import GoogleAuth
from pydrive.drive import GoogleDrive

headers = ["RSVP for Convocation AGM", "07 December 2019"]
table_headers = ["First names", "Surname", "ID No", "Cell Nlo", "Email Address"]
document = Document()
document.add_heading(headers[0])
document.add_heading(headers[1], level=2)

table = document.add_table(rows=10, cols=5)

index = 0
heading_cells = table.rows[0].cells
for header_name in table_headers:
    heading_cells[index].text = header_name
    index += 1

items = (
    (7, '1024', 'Plush kittens', "asdasd", "663242388"),
    (3, '2042', 'Furbees', "asdasd", "663242388"),
    (1, '1288', 'French Poodle Collars, Deluxe', "asdasd", "663242388"),
)
for item in items:
    cells = table.add_row().cells
    cells[0].text = str(item[0])
    cells[1].text = item[1]
    cells[2].text = item[2]
    cells[3].text = item[3]
    cells[4].text = item[4]

    
    

document.save('rsvp_table.docx')